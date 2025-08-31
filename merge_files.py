from docx import Document
from docx.shared import Inches, Pt
from pathlib import Path

# Paths
cards_dir = Path("./Employee_Reports23/TrainingIDs")
output_docx = Path("merged_idCards.docx")

# Create document
doc = Document()
doc.add_heading("Employee Training ID Cards", level=1)

# Get all image paths sorted
image_paths = sorted(cards_dir.glob("ID_CARD_*.png"))

# Create a table for 2 ID cards per row
table = doc.add_table(rows=0, cols=2)
table.autofit = False

# Adjust table width & spacing between columns
for i in range(0, len(image_paths), 2):
    row_cells = table.add_row().cells
    
    # Insert first ID card
    p1 = row_cells[0].paragraphs[0]
    run1 = p1.add_run()
    run1.add_picture(str(image_paths[i]), width=Inches(3.0))  # Reduced width to fit page
    p1.alignment = 1  # Center alignment
    
    # Insert second ID card if available
    if i + 1 < len(image_paths):
        p2 = row_cells[1].paragraphs[0]
        run2 = p2.add_run()
        run2.add_picture(str(image_paths[i + 1]), width=Inches(3.0))
        p2.alignment = 1  # Center alignment

# Add spacing between rows (vertical padding)
for row in table.rows:
    for cell in row.cells:
        # Set bottom padding of 10px (~0.15 inches)
        cell_paragraph = cell.paragraphs[0]
        cell_paragraph.space_after = Pt(10)

# Save the document
doc.save(output_docx)
print(f"Saved to {output_docx}")
